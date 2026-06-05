import json
import streamlit as st
import requests
from bs4 import BeautifulSoup
import google.generativeai as genai
import datetime
import io
from docx import Document

# ==========================================
# 1. ページ初期設定＆デザインシステム (CSS)
# ==========================================
st.set_page_config(
    page_title="医療広告ガイドラインチェッカー",
    page_icon="🏥",
    layout="wide",
    initial_sidebar_state="collapsed",
)

st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Noto+Sans+JP:wght@300;400;500;700&display=swap');
    html, body, [data-testid="stAppViewContainer"] {
        font-family: 'Noto Sans JP', sans-serif;
    }

    .header-container {
        padding: 2.5rem 1.5rem;
        background: linear-gradient(135deg, #1e3a8a 0%, #3b82f6 100%);
        color: white;
        border-radius: 16px;
        margin-bottom: 2rem;
        box-shadow: 0 4px 20px rgba(0, 0, 0, 0.1);
        text-align: center;
    }
    .header-title {
        font-size: 2.5rem;
        font-weight: 700;
        margin-bottom: 0.5rem;
        letter-spacing: -0.02em;
    }
    .header-subtitle {
        font-size: 1.1rem;
        opacity: 0.9;
        font-weight: 300;
    }

    .status-box {
        padding: 1.5rem;
        border-radius: 12px;
        font-weight: bold;
        text-align: center;
        font-size: 1.3rem;
        margin-bottom: 1.5rem;
        box-shadow: 0 4px 6px rgba(0,0,0,0.05);
    }
    .status-ng {
        background-color: #fef2f2;
        color: #dc2626;
        border: 1px solid #fee2e2;
    }
    .status-warning {
        background-color: #fffbeb;
        color: #d97706;
        border: 1px solid #fef3c7;
    }
    .status-ok {
        background-color: #f0fdf4;
        color: #16a34a;
        border: 1px solid #dcfce7;
    }

    .card {
        background-color: white;
        color: #1f2937;
        padding: 1.5rem;
        border-radius: 12px;
        border-left: 6px solid #e5e7eb;
        margin-bottom: 1.25rem;
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.05);
        transition: transform 0.2s ease, box-shadow 0.2s ease;
    }
    .card:hover {
        transform: translateY(-2px);
        box-shadow: 0 6px 12px rgba(0, 0, 0, 0.08);
    }
    .card-ng { border-left-color: #dc2626; }
    .card-warning { border-left-color: #d97706; }

    .badge {
        display: inline-block;
        padding: 0.25rem 0.6rem;
        font-size: 0.75rem;
        font-weight: 600;
        border-radius: 9999px;
        margin-right: 0.5rem;
        text-transform: uppercase;
    }
    .badge-high { background-color: #fee2e2; color: #dc2626; }
    .badge-mid  { background-color: #fef3c7; color: #d97706; }

    .original-phrase {
        font-size: 1.1rem;
        font-weight: 700;
        color: #111827;
        margin-bottom: 0.5rem;
        background-color: #f3f4f6;
        padding: 0.4rem 0.8rem;
        border-radius: 6px;
        display: inline-block;
    }
    .meta-section {
        font-size: 0.85rem;
        color: #4b5563;
        margin-top: 0.5rem;
        padding-top: 0.5rem;
        border-top: 1px dashed #e5e7eb;
    }
</style>
""", unsafe_allow_html=True)

st.markdown("""
<div class="header-container">
    <div class="header-title">🏥 医療広告ガイドラインチェッカー</div>
    <div class="header-subtitle">AIによる3段階検証で、医療広告規制への抵触リスクと改善案の品質を可視化します</div>
</div>
""", unsafe_allow_html=True)

# ==========================================
# 2. ヘルパー関数
# ==========================================

def extract_text_from_url(url: str) -> str:
    """URLからメインテキストをスクレイピングする"""
    headers = {
        "User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
    }
    try:
        response = requests.get(url, headers=headers, timeout=10)
        response.raise_for_status()
        response.encoding = response.apparent_encoding
        soup = BeautifulSoup(response.text, 'html.parser')
        for tag in soup(["script", "style", "nav", "footer", "header", "iframe"]):
            tag.decompose()
        lines = (line.strip() for line in soup.get_text().splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = "\n".join(chunk for chunk in chunks if chunk)
        if not text or len(text.strip()) < 10:
            raise ValueError("抽出されたテキストが極めて少ないか、空です。")
        return text
    except Exception as e:
        raise RuntimeError(f"URLからのデータ取得に失敗しました: {str(e)}")


def run_gemini_stage_1(text: str, model_name: str) -> str:
    """ステップ1: 違反の疑いがある箇所を一次抽出する"""
    prompt = f"""
あなたは医療広告ガイドラインの監査専門家です。
以下の医療広告原稿（または医療Webサイトのテキスト）を厳格に分析し、
「医療広告ガイドライン」「医療法」に抵触するリスクがある箇所をリストアップしてください。

【分析観点】
1. 虚偽広告: 「絶対安全」「100%治る」等
2. 誇大広告: 「地域No.1」「最新の治療法」等
3. 比較優良広告: 「他院より優れている」「日本一」等
4. 主観的な体験談: 患者の主観的な感想、治療成果の体験談等
5. 客観的証明が困難な内容: データや根拠のない効能表現等
6. 品位を損ねる内容: 「今だけ半額」「格安キャンペーン」など費用を過度に強調する表現等
7. 無関係な事項: 「著名人〇〇様も来院」などの誘引表現等

【出力フォーマット】
以下のキーを持つJSONの配列として出力してください。Markdownのjsonコードブロック（```json ... ```）で囲んでください。
[
  {{
    "phrase": "抵触リスクのある具体的なフレーズ",
    "risk_level": "高" または "中",
    "category": "違反カテゴリ（誇大広告、主観的な体験談など）",
    "reason": "ガイドラインに抵触する理由の解説",
    "legal_basis": "抵触する恐れのあるガイドライン項目や条文名",
    "legal_basis_date": "そのガイドライン・条文の発出・最終改訂年月日（例: 平成30年5月8日改正、など。不確かな場合は空欄）",
    "suggestion": "具体的な代替表現または削除の推奨案"
  }}
]

【分析対象テキスト】
{text}
"""
    model = genai.GenerativeModel(model_name)
    response = model.generate_content(prompt)
    return response.text


def run_gemini_stage_2(stage1_result: str, model_name: str) -> str:
    """ステップ2: 法的根拠のファクトチェック・ダブルチェック"""
    prompt = f"""
あなたは医療法および医療広告ガイドラインに極めて精通した法律の専門家・ファクトチェッカーです。
ステップ1でAIが一次抽出した「ガイドライン違反の疑いリスト」に対して、厳格なファクトチェックを行ってください。

【ファクトチェックのルール】
1. 「法的根拠 (legal_basis)」および「発出・最終改訂年月日 (legal_basis_date)」が、実在する公式の医療広告ガイドラインや厚生労働省の告示・通達と一致しているかをファクトチェックしてください。
2. 存在しない架空の法律・条文や、誤った年月日（ハルシネーション）が書かれている場合は、正しい情報に修正してください。
3. もし「これは明らかにガイドラインに抵触しない」「法的な根拠が全く実在しない誤認である」と判断される指摘があった場合は、その指摘項目そのものをリストから完全に削除（除外）してください。
4. 指摘内容自体が実在のガイドライン項目に準拠している場合は、法的根拠と改訂日を正確に補正した上で残してください。

【出力フォーマット】
最終的な精査済みのリストを、以下のキーを持つJSONの配列として出力してください。Markdownのjsonコードブロック（```json ... ```）で囲んでください。
[
  {{
    "phrase": "修正済みの抵触リスクのあるフレーズ",
    "risk_level": "高" または "中",
    "category": "違反カテゴリ",
    "reason": "ファクトチェックを経た正確な理由の解説",
    "legal_basis": "正確に実在するガイドライン項目や条文名",
    "legal_basis_date": "実在する発出・最終改訂年月日（例: 平成30年5月8日改正、令和5年4月1日一部改正など。判明している最新の日付を正確に記載）",
    "suggestion": "具体的な代替表現または削除の推奨案"
  }}
]

【検証対象（ステップ1の抽出結果）】
{stage1_result}
"""
    model = genai.GenerativeModel(model_name)
    response = model.generate_content(prompt)
    return response.text


def run_gemini_stage_3(stage2_result: str, model_name: str) -> str:
    """ステップ3: 改善案の日本語品質チェック＆ガイドライン最終確認"""
    prompt = f"""
あなたは医療広告ガイドラインの専門家であり、かつ日本語表現の校正・編集者でもあります。
ステップ2までの精査を経たリストに対して、各指摘の「改善案 (suggestion)」について以下の2点を最終確認・修正してください。

【確認・修正のルール】
1. 【日本語品質チェック】
   - 「改善案」の文章が、日本語として正確かつ自然な表現になっているかを確認してください。
   - 「てにをは」の誤り、不自然な言い回し、冗長な表現、二重否定、主語と述語のねじれなどがあれば修正してください。
   - 実際の広告・Webサイトで使用できる、洗練された自然な日本語に整えてください。

2. 【ガイドライン最終確認】
   - 修正後の「改善案」が、医療広告ガイドラインおよび医療法に抵触しない安全な表現であることを最終確認してください。
   - もし改善案自体がまだガイドラインに抵触する恐れがある場合は、安全な表現に修正してください。
   - 改善案が「削除を推奨する」という趣旨の場合は、その旨を簡潔かつ明確な日本語で表現してください。

【出力フォーマット】
最終的な完成版リストを、以下のキーを持つJSONの配列として出力してください。Markdownのjsonコードブロック（```json ... ```）で囲んでください。
[
  {{
    "phrase": "抵触リスクのあるフレーズ",
    "risk_level": "高" または "中",
    "category": "違反カテゴリ",
    "reason": "理由の解説",
    "legal_basis": "実在するガイドライン項目や条文名",
    "legal_basis_date": "発出・最終改訂年月日",
    "suggestion": "日本語として自然かつガイドラインに完全準拠した改善案"
  }}
]

【検証対象（ステップ2の精査済み結果）】
{stage2_result}
"""
    model = genai.GenerativeModel(model_name)
    response = model.generate_content(prompt)
    return response.text


def parse_json_from_gemini(raw_text: str):
    """Geminiの出力からJSONブロックを抽出してデコードする"""
    text = raw_text.strip()
    if "```json" in text:
        text = text.split("```json")[1].split("```")[0].strip()
    elif "```" in text:
        text = text.split("```")[1].split("```")[0].strip()
    return json.loads(text)


def generate_docx(results: list) -> bytes:
    """分析結果をWord（.docx）形式のバイナリデータとして生成する"""
    doc = Document()
    doc.add_heading("医療広告ガイドライン分析結果", level=1)
    high_risk_count = sum(1 for r in results if r.get("risk_level") == "高")
    doc.add_paragraph(f"総指摘件数: {len(results)} 件（うち 危険度「高」: {high_risk_count} 件）\n")
    for idx, res in enumerate(results, 1):
        doc.add_heading(f"指摘 {idx}", level=2)
        p = doc.add_paragraph()
        p.add_run("対象表現: ").bold = True
        p.add_run(f"「{res.get('phrase', '')}」\n")
        p.add_run("危険度: ").bold = True
        p.add_run(f"{res.get('risk_level', '')}\n")
        p.add_run("カテゴリ: ").bold = True
        p.add_run(f"{res.get('category', '')}\n")
        p.add_run("抵触理由: ").bold = True
        p.add_run(f"{res.get('reason', '')}\n")
        p.add_run("改善案（代替表現・削除案）: ").bold = True
        p.add_run(f"{res.get('suggestion', '')}\n")
        p.add_run("法的根拠: ").bold = True
        p.add_run(f"{res.get('legal_basis', '未詳')} (発出・最終改訂: {res.get('legal_basis_date', '不明・未記載')})")
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


def generate_txt(results: list) -> str:
    """分析結果をプレーンテキスト形式として生成する"""
    lines = []
    lines.append("=========================================")
    lines.append("医療広告ガイドライン 分析結果")
    lines.append("=========================================\n")
    high_risk_count = sum(1 for r in results if r.get("risk_level") == "高")
    lines.append(f"総指摘件数: {len(results)} 件（うち 危険度「高」: {high_risk_count} 件）\n")
    for idx, res in enumerate(results, 1):
        lines.append(f"--- 指摘 {idx} ---")
        lines.append(f"対象表現: 「{res.get('phrase', '')}」")
        lines.append(f"危険度: {res.get('risk_level', '')}")
        lines.append(f"カテゴリ: {res.get('category', '')}")
        lines.append(f"抵触理由: {res.get('reason', '')}")
        lines.append(f"改善案（代替表現・削除案）: {res.get('suggestion', '')}")
        lines.append(f"法的根拠: {res.get('legal_basis', '未詳')} (発出・最終改訂: {res.get('legal_basis_date', '不明・未記載')})\n")
    return "\n".join(lines)


# ==========================================
# 3. アプリケーションUI & メインロジック
# ==========================================

api_key = None
if "GEMINI_API_KEY" in st.secrets:
    api_key = st.secrets["GEMINI_API_KEY"]

if not api_key or api_key == "YOUR_GEMINI_API_KEY_HERE":
    st.warning("⚠️ Gemini APIキーが設定されていません。ローカル環境の場合は `.streamlit/secrets.toml` に `GEMINI_API_KEY = \"あなたのキー\"` を記述するか、デプロイ環境の環境変数・Secretsに設定してください。")
    temp_key = st.text_input("一時的に使用するGemini APIキーを入力することも可能です:", type="password")
    if temp_key:
        api_key = temp_key

if api_key and api_key != "YOUR_GEMINI_API_KEY_HERE":
    genai.configure(api_key=api_key)

    with st.sidebar:
        st.header("⚙️ 設定")
        model_name = st.selectbox(
            "AIモデルの選択",
            ["gemini-2.5-flash", "gemini-1.5-pro", "gemini-1.5-flash"],
            index=0,
            help="通常は高速で高性能な gemini-2.5-flash で問題ありません。より緻密な判断を求める場合は gemini-1.5-pro を選択してください。"
        )
        st.markdown("---")
        st.markdown("""
        **【医療広告ガイドラインの主な監視項目】**
        - 虚偽広告 / 誇大広告
        - 比較優良広告
        - 患者の主観的体験談
        - 客観的証明が困難な内容
        - 品位を損ねる内容（不適切な割引等）
        - 医療と無関係な誘引（著名人の来院等）

        **【3段階チェックの流れ】**
        1. 一次抽出（違反候補の特定）
        2. 法的根拠のファクトチェック
        3. 改善案の日本語校正＋ガイドライン最終確認
        """)

    st.subheader("📝 広告原稿・Webページの入力")
    input_method = st.radio("入力方法を選択してください:", ["テキストを直接入力", "URLから読み込み"], horizontal=True)

    input_text = ""
    url_input = ""

    if input_method == "テキストを直接入力":
        input_text = st.text_area(
            "チェックしたい広告原稿・テキストを貼り付けてください（例: チラシ原稿、LPの文言など）",
            height=250,
            placeholder="（例）\n当院は地域No.1の審美歯科クリニックです！\n絶対に失敗しない最新のインプラント治療を、今なら半額キャンペーン実施中！\nモデルの〇〇さんも「劇的に痛みが消えた！」と大絶賛しています。"
        )
    else:
        url_input = st.text_input("チェックしたいWebページのURLを入力してください:", placeholder="https://example.com/clinic-page")
        if url_input:
            with st.spinner("🔗 URLからテキストを抽出中..."):
                try:
                    input_text = extract_text_from_url(url_input)
                    st.success("テキストの抽出に成功しました！以下の抽出内容を確認して「チェック実行」を押してください。")
                    with st.expander("📄 抽出されたテキストを確認する"):
                        st.text_area("抽出テキスト（編集可能）", value=input_text, height=200, key="scraped_text_area")
                        if "scraped_text_area" in st.session_state:
                            input_text = st.session_state["scraped_text_area"]
                except Exception as e:
                    st.error(f"❌ URLからコンテンツを読み込めませんでした。\n理由: {str(e)}")
                    st.info("💡 Bot防止機能などが働いている可能性があります。恐れ入りますが、「テキストを直接入力」タブに切り替え、Webサイトの文章をコピー＆ペーストしてご利用ください。")

    # セッションステートの初期化
    if "analysis_results" not in st.session_state:
        st.session_state.analysis_results = None
    if "checked_text" not in st.session_state:
        st.session_state.checked_text = ""

    if input_text:
        # 入力テキストが変更された場合は古い結果をクリア
        if input_text != st.session_state.checked_text and st.session_state.analysis_results is not None:
            st.session_state.analysis_results = None
            st.session_state.checked_text = ""

        if st.button("🔍 ガイドラインチェックを実行する", type="primary", use_container_width=True):
            st.session_state.analysis_results = None
            st.session_state.checked_text = ""

            progress_bar = st.progress(0)
            status_text = st.empty()

            try:
                # ステップ1: 一次抽出
                status_text.markdown("🔄 **[Step 1/3]** AIがガイドライン抵触の疑いがある表現を一次抽出しています...")
                progress_bar.progress(15)
                stage1_raw = run_gemini_stage_1(input_text, model_name)
                progress_bar.progress(35)

                # ステップ2: 法的根拠のファクトチェック
                status_text.markdown("🛡️ **[Step 2/3]** 法的根拠・条文・発出日の正確性をファクトチェックしています...")
                progress_bar.progress(50)
                stage2_raw = run_gemini_stage_2(stage1_raw, model_name)
                progress_bar.progress(70)

                # ステップ3: 改善案の日本語品質＆ガイドライン最終確認
                status_text.markdown("✍️ **[Step 3/3]** 改善案の日本語品質とガイドライン準拠を最終確認しています...")
                progress_bar.progress(85)
                stage3_raw = run_gemini_stage_3(stage2_raw, model_name)
                progress_bar.progress(95)

                status_text.markdown("📊 結果を整形中...")
                results = parse_json_from_gemini(stage3_raw)

                progress_bar.progress(100)
                status_text.empty()
                progress_bar.empty()

                st.session_state.analysis_results = results
                st.session_state.checked_text = input_text

            except json.JSONDecodeError:
                status_text.empty()
                progress_bar.empty()
                st.error("❌ AIの解析結果を正しく処理できませんでした。出力フォーマットが崩れた可能性があります。再度実行してください。")
                if 'stage3_raw' in locals():
                    with st.expander("生データを確認する"):
                        st.text(stage3_raw)
            except Exception as e:
                status_text.empty()
                progress_bar.empty()
                st.error(f"❌ エラーが発生しました: {str(e)}")
    else:
        st.info("💡 テキストを入力するか、URLを読み込んで「ガイドラインチェックを実行する」ボタンを押してください。")

    # 分析結果の描画
    if st.session_state.analysis_results is not None:
        results = st.session_state.analysis_results

        st.markdown("---")
        st.subheader("📊 分析結果")

        if not results:
            st.markdown("""
            <div class="status-box status-ok">
                ✅ 医療広告ガイドラインに抵触する可能性のある表現は見つかりませんでした。
            </div>
            """, unsafe_allow_html=True)
            st.balloons()
        else:
            high_risk_count = sum(1 for r in results if r.get("risk_level") == "高")

            if high_risk_count > 0:
                st.markdown(f"""
                <div class="status-box status-ng">
                    ❌ 要注意・修正必須の表現が {len(results)} 件検出されました。（うち 危険度「高」: {high_risk_count} 件）
                </div>
                """, unsafe_allow_html=True)
            else:
                st.markdown(f"""
                <div class="status-box status-warning">
                    ⚠️ 要注意表現が {len(results)} 件検出されました。念のため内容の見直しを推奨します。
                </div>
                """, unsafe_allow_html=True)

            for idx, res in enumerate(results, 1):
                risk = res.get("risk_level", "中")
                risk_class = "card-ng" if risk == "高" else "card-warning"
                badge_class = "badge-high" if risk == "高" else "badge-mid"

                st.markdown(f"""
                <div class="card {risk_class}">
                    <div style="display: flex; align-items: center; margin-bottom: 0.75rem;">
                        <span class="badge {badge_class}">危険度: {risk}</span>
                        <span style="font-weight: 600; color: #374151; font-size: 0.95rem;">カテゴリ: {res.get('category', '一般内規')}</span>
                    </div>
                    <div class="original-phrase">対象表現: 「{res.get('phrase', '')}」</div>
                    <div style="margin-top: 0.75rem; font-size: 0.95rem; line-height: 1.6;">
                        <strong>💡 抵触理由:</strong> {res.get('reason', '')}
                    </div>
                    <div style="margin-top: 0.5rem; font-size: 0.95rem; line-height: 1.6; background-color: #f0fdf4; padding: 0.5rem 0.8rem; border-radius: 6px; border-left: 3px solid #16a34a;">
                        <strong>✅ 改善案（代替表現・削除案）:</strong> {res.get('suggestion', '')}
                    </div>
                    <div class="meta-section">
                        ⚖️ <strong>法的根拠:</strong> {res.get('legal_basis', '未詳')}
                        <span style="margin-left: 1.5rem;">📅 <strong>発出・最終改訂:</strong> {res.get('legal_basis_date', '不明・未記載')}</span>
                    </div>
                </div>
                """, unsafe_allow_html=True)

        # ダウンロードボタン（Word形式・テキスト形式）
        st.markdown("---")
        col1, col2 = st.columns(2)
        with col1:
            docx_data = generate_docx(results)
            st.download_button(
                label="📥 Word形式でダウンロード (.docx)",
                data=docx_data,
                file_name="analysis_results.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
        with col2:
            txt_data = generate_txt(results)
            st.download_button(
                label="📥 テキスト形式でダウンロード (.txt)",
                data=txt_data,
                file_name="analysis_results.txt",
                mime="text/plain",
                use_container_width=True
            )

else:
    st.info("APIキーの入力を待機しています。")
